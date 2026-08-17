"""Word document parser (python-docx)."""

from __future__ import annotations

from pathlib import Path

from backend.core.models import ParsedDocument, TableResult
from backend.ingestion.text_parser import FileParseError
from backend.utils.logger import logger

# Word has no fixed pages until rendered — approximate for citation purposes.
PARAGRAPHS_PER_PAGE = 28


class DOCXParser:
    SUPPORTED = {".docx"}

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        try:
            import docx
        except ImportError as exc:  # pragma: no cover - declared dependency
            raise FileParseError(f"python-docx not installed: {exc}") from exc

        try:
            document = docx.Document(str(path))
        except Exception as exc:
            raise FileParseError(f"Cannot open {path.name}: {exc}") from exc

        warnings: list[str] = []
        parts: list[str] = []
        headings: list[str] = []
        page = 1
        paragraph_count = 0
        parts.append(f"\n\n--- PAGE {page} ---\n")

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style = (paragraph.style.name if paragraph.style else "") or ""
            if style.lower().startswith("heading"):
                headings.append(text)
                parts.append(f"\n\n{text.upper()}\n")
            else:
                parts.append(text + "\n")

            paragraph_count += 1
            if paragraph_count % PARAGRAPHS_PER_PAGE == 0:
                page += 1
                parts.append(f"\n\n--- PAGE {page} ---\n")

        tables: list[TableResult] = []
        for index, table in enumerate(document.tables):
            try:
                markdown = _table_to_markdown(table)
            except Exception as exc:
                warnings.append(f"Table {index + 1} could not be parsed: {exc}")
                continue
            rows = len(table.rows)
            cols = len(table.columns) if rows else 0
            tables.append(TableResult(page=page, markdown=markdown, rows=rows, cols=cols))
            parts.append(f"\n\n[TABLE {index + 1}]\n{markdown}\n")

        text = "".join(parts)
        if not text.strip():
            raise FileParseError(f"{path.name} contains no extractable text")

        core = document.core_properties
        logger.info(
            f"Parsed docx {path.name}: {len(text)} chars, ~{page} pages, "
            f"{len(tables)} tables"
        )
        return ParsedDocument(
            text=text,
            file_name=path.name,
            page_count=page,
            metadata={
                "title": (core.title or path.stem) if core else path.stem,
                "author": (core.author or "") if core else "",
                "creation_date": str(core.created) if core and core.created else "",
                "source_format": "docx",
                "heading_count": len(headings),
                "section_names": headings[:60],
                "table_count": len(tables),
            },
            tables=tables,
            warnings=warnings,
        )


def _table_to_markdown(table) -> str:
    rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    header = rows[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return "\n".join(lines)
